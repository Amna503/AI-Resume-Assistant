import json
import os
import re
import tempfile
from pathlib import Path

import streamlit as st
from google import genai
from pypdf import PdfReader
from docx import Document


# -----------------------------
# App configuration
# -----------------------------
st.set_page_config(
    page_title="Resume ATS Analyzer",
    page_icon="📄",
    layout="wide",
)

MODEL_NAME = "gemini-2.5-flash"
MAX_FILE_SIZE_MB = 10


# -----------------------------
# Styling
# -----------------------------
st.markdown(
    """
    <style>
        .main-title {
            font-size: 2.4rem;
            font-weight: 700;
            margin-bottom: 0.2rem;
        }
        .subtitle {
            color: #6b7280;
            margin-bottom: 1.5rem;
        }
        .score-card {
            padding: 1.2rem;
            border: 1px solid rgba(128,128,128,.25);
            border-radius: 14px;
            text-align: center;
        }
        .score {
            font-size: 3rem;
            font-weight: 800;
            line-height: 1;
        }
        .small-note {
            color: #6b7280;
            font-size: .9rem;
        }
    </style>
    """,
    unsafe_allow_html=True,
)


# -----------------------------
# Helpers
# -----------------------------
def get_api_key():
    """Read Gemini API key from Streamlit secrets first, then environment."""
    try:
        if "GEMINI_API_KEY" in st.secrets:
            return st.secrets["GEMINI_API_KEY"]
    except Exception:
        pass

    return os.getenv("GEMINI_API_KEY")


def extract_pdf_text(file_bytes):
    """Extract selectable text from a PDF."""
    reader = PdfReader(__import__("io").BytesIO(file_bytes))
    pages = []

    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text)

    return "\n\n".join(pages).strip()


def extract_docx_text(file_bytes):
    """Extract text from paragraphs and tables in a DOCX file."""
    from io import BytesIO

    doc = Document(BytesIO(file_bytes))
    chunks = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            chunks.append(paragraph.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                chunks.append(row_text)

    return "\n".join(chunks).strip()


def extract_text(uploaded_file):
    """Return extracted text and whether Gemini should inspect the original file."""
    file_bytes = uploaded_file.getvalue()
    suffix = Path(uploaded_file.name).suffix.lower()

    if suffix == ".pdf":
        text = extract_pdf_text(file_bytes)
    elif suffix == ".docx":
        text = extract_docx_text(file_bytes)
    else:
        raise ValueError("Unsupported file type.")

    # Very little extracted text usually means a scanned/image-only PDF.
    needs_original_file = len(text.strip()) < 200
    return text, needs_original_file, file_bytes


def clean_json_response(raw_text):
    """Extract JSON even if Gemini wraps it in Markdown fences."""
    text = (raw_text or "").strip()
    text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s*```$", "", text)

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", text, flags=re.DOTALL)
        if not match:
            raise ValueError("Gemini returned an invalid JSON response.")
        return json.loads(match.group(0))


def clamp_score(value):
    try:
        return max(0, min(100, int(float(value))))
    except (TypeError, ValueError):
        return 0


def analyze_resume(client, resume_text, job_description="", original_file=None):
    """
    Analyze a resume with Gemini and return a structured dictionary.

    If resume_text is too short, the original PDF/DOCX is uploaded to Gemini
    so the model can inspect the document directly.
    """
    job_context = (
        job_description.strip()
        if job_description.strip()
        else "No specific job description was provided. Evaluate general ATS readiness."
    )

    prompt = f"""
You are an expert ATS resume evaluator and professional resume coach.

Analyze the candidate's resume and return ONLY valid JSON.

IMPORTANT:
- Do not invent experience, skills, education, employers, dates, certifications, or achievements.
- Base every recommendation on the resume and, when supplied, the job description.
- ATS score is a practical heuristic, not a score from a specific ATS vendor.
- If no job description is supplied, evaluate general ATS readiness.
- Prefer concrete, actionable improvements.
- A strong resume should use standard section headings, clear chronology,
  measurable achievements, relevant keywords, simple formatting, and concise language.

JOB DESCRIPTION:
{job_context}

Return exactly this JSON structure:
{{
  "ats_score": 0,
  "score_breakdown": {{
    "keyword_match": 0,
    "formatting_ats_readiness": 0,
    "experience_impact": 0,
    "sections_completeness": 0,
    "clarity_readability": 0
  }},
  "summary": "2-4 sentence overall assessment",
  "strengths": [
    "strength 1",
    "strength 2",
    "strength 3"
  ],
  "critical_improvements": [
    {{
      "issue": "specific issue",
      "why_it_matters": "ATS/recruiter impact",
      "how_to_fix": "specific fix"
    }}
  ],
  "keyword_suggestions": [
    "keyword or phrase to consider"
  ],
  "formatting_checks": [
    "specific formatting observation"
  ],
  "content_improvements": [
    {{
      "section": "Experience",
      "before": "short excerpt or description from the resume, if available",
      "after": "improved example that does not add unsupported facts"
    }}
  ],
  "action_plan": [
    "highest priority action",
    "second priority action",
    "third priority action"
  ]
}}

Scoring guidance:
- keyword_match: 0-30 points
- formatting_ats_readiness: 0-20 points
- experience_impact: 0-20 points
- sections_completeness: 0-15 points
- clarity_readability: 0-15 points
- ats_score must equal the sum of the five categories and stay between 0 and 100.

Resume text:
{resume_text[:60000]}
"""

    contents = [prompt]

    temp_path = None
    if original_file is not None:
        suffix = ".pdf" if original_file[:4] == b"%PDF" else ".docx"
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(original_file)
            temp_path = tmp.name

        try:
            uploaded = client.files.upload(file=temp_path)
            contents = [prompt, uploaded]
        finally:
            try:
                os.remove(temp_path)
            except OSError:
                pass

    response = client.models.generate_content(
        model=MODEL_NAME,
        contents=contents,
    )

    return clean_json_response(response.text)


def display_bullets(items):
    if not items:
        st.write("No items returned.")
        return

    for item in items:
        if isinstance(item, dict):
            title = item.get("issue", "Improvement")
            why = item.get("why_it_matters", "")
            fix = item.get("how_to_fix", "")
            st.markdown(f"**{title}**")
            if why:
                st.write(f"Why it matters: {why}")
            if fix:
                st.write(f"How to fix: {fix}")
        else:
            st.markdown(f"- {item}")


# -----------------------------
# UI
# -----------------------------
st.markdown('<div class="main-title">📄 Resume ATS Analyzer</div>', unsafe_allow_html=True)
st.markdown(
    '<div class="subtitle">Upload your resume to get an AI-powered ATS readiness score, '
    'keyword analysis, and practical improvements.</div>',
    unsafe_allow_html=True,
)

with st.sidebar:
    st.header("Settings")
    st.markdown(
        "Supported files: **PDF** and **DOCX**\n\n"
        f"Maximum file size: **{MAX_FILE_SIZE_MB} MB**"
    )
    st.divider()
    st.markdown(
        "**How the score works**\n\n"
        "The AI evaluates keyword relevance, ATS-friendly formatting, "
        "achievement impact, section completeness, and readability."
    )
    st.caption(f"Model: {MODEL_NAME}")

api_key = get_api_key()

if not api_key:
    st.error(
        "Gemini API key not found. Add `GEMINI_API_KEY` to "
        "`.streamlit/secrets.toml` locally or to Streamlit Cloud Secrets."
    )
    st.stop()

uploaded_file = st.file_uploader(
    "Upload your resume",
    type=["pdf", "docx"],
    help="Use a text-based PDF or DOCX for the best result.",
)

job_description = st.text_area(
    "Optional: paste the job description",
    height=180,
    placeholder="Adding the target job description makes the keyword/ATS analysis more specific.",
)

analyze_button = st.button(
    "🔍 Analyze Resume",
    type="primary",
    use_container_width=True,
)

if analyze_button:
    if uploaded_file is None:
        st.warning("Please upload a PDF or DOCX resume first.")
        st.stop()

    if uploaded_file.size > MAX_FILE_SIZE_MB * 1024 * 1024:
        st.error(f"File is too large. Please upload a file under {MAX_FILE_SIZE_MB} MB.")
        st.stop()

    try:
        resume_text, needs_original_file, file_bytes = extract_text(uploaded_file)

        if not resume_text and not needs_original_file:
            st.error("Could not extract text from this resume.")
            st.stop()

        with st.spinner("Analyzing your resume with Gemini Flash..."):
            client = genai.Client(api_key=api_key)

            # For normal text PDFs/DOCX, text extraction keeps the request small.
            # For scanned/image-heavy PDFs, send the original file to Gemini.
            original_for_gemini = file_bytes if needs_original_file else None

            result = analyze_resume(
                client=client,
                resume_text=resume_text,
                job_description=job_description,
                original_file=original_for_gemini,
            )

        ats_score = clamp_score(result.get("ats_score"))
        st.session_state["analysis"] = result
        st.session_state["ats_score"] = ats_score

    except Exception as exc:
        st.error(f"Analysis failed: {exc}")
        st.info(
            "Check your Gemini API key, internet connection, file type, "
            "and Gemini API quota. If the resume is scanned, try a clearer PDF."
        )


# -----------------------------
# Results
# -----------------------------
if "analysis" in st.session_state:
    result = st.session_state["analysis"]
    ats_score = st.session_state["ats_score"]

    st.divider()

    col1, col2 = st.columns([1, 2])

    with col1:
        st.markdown(
            f"""
            <div class="score-card">
                <div class="small-note">ATS Readiness Score</div>
                <div class="score">{ats_score}/100</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    with col2:
        st.subheader("Overall assessment")
        st.write(result.get("summary", "No summary available."))

    st.subheader("📊 Score breakdown")

    breakdown = result.get("score_breakdown", {})
    score_cols = st.columns(5)
    labels = [
        ("Keyword Match", "keyword_match"),
        ("ATS Formatting", "formatting_ats_readiness"),
        ("Experience Impact", "experience_impact"),
        ("Sections", "sections_completeness"),
        ("Clarity", "clarity_readability"),
    ]

    max_scores = {
        "keyword_match": 30,
        "formatting_ats_readiness": 20,
        "experience_impact": 20,
        "sections_completeness": 15,
        "clarity_readability": 15,
    }

    for col, (label, key) in zip(score_cols, labels):
        value = clamp_score(breakdown.get(key))
        maximum = max_scores[key]
        with col:
            st.metric(label, f"{value}/{maximum}")

    tab1, tab2, tab3, tab4, tab5 = st.tabs(
        [
            "✅ Strengths",
            "🚨 Improvements",
            "🔑 Keywords",
            "📝 Content",
            "🎯 Action Plan",
        ]
    )

    with tab1:
        display_bullets(result.get("strengths", []))

    with tab2:
        display_bullets(result.get("critical_improvements", []))
        st.subheader("Formatting checks")
        display_bullets(result.get("formatting_checks", []))

    with tab3:
        keywords = result.get("keyword_suggestions", [])
        if keywords:
            st.write("Consider these keywords/phrases where they are genuinely relevant:")
            st.code(", ".join(str(k) for k in keywords), language="text")
        else:
            st.write("No keyword suggestions returned.")

    with tab4:
        improvements = result.get("content_improvements", [])
        if not improvements:
            st.write("No content rewrites returned.")
        else:
            for item in improvements:
                st.markdown(f"### {item.get('section', 'Section')}")
                st.markdown("**Current:**")
                st.write(item.get("before", ""))
                st.markdown("**Suggested:**")
                st.write(item.get("after", ""))

    with tab5:
        display_bullets(result.get("action_plan", []))

    st.caption(
        "Note: This is an AI-based ATS readiness estimate, not a guarantee of how a "
        "specific employer's ATS will score the resume."
    )
